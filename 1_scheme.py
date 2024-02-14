import math2docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
from io import BytesIO
import sympy as sp
import re
import cmath

doc = Document("new.docx")

# Page №3

text = doc.add_paragraph("Для заданной схемы дано:")
text.runs[0].font.size = Pt(14)
text.paragraph_format.left_indent = Mm(15)
text = doc.add_paragraph("Таблица 1")
text.alignment = WD_ALIGN_PARAGRAPH.RIGHT

start_table1 = {'0': [100, 1, -120], '1': [300, 5.5, 90], '2': [280, 5, 60], '3': [260, 4.5, 45],
                '4': [240, 4, 30], '5': [220, 3.5, 0], '6': [200, 3, -30], '7': [175, 2.5, -45],
                '8': [150, 2, -60], '9': [125, 1.5, -90]}
start_table2 = {'0': [1000, 10, 0.02, 200], '1': [100, 100, 2, 200], '2': [150, 90, 1.2, 150],
                '3': [200, 80, 0.8, 125], '4': [250, 75, 0.6, 107], '5': [300, 60, 0.4, 111],
                '6': [400, 50, 0.25, 100], '7': [500, 40, 0.16, 100],
                '8': [600, 30, 0.1, 111], '9': [800, 24, 0.06, 200]}
# Получаем ввод от пользователя
input_row1 = int(input("Введите номер строки для таблицы 1: "))
input_row2 = int(input("Введите номер строки для таблицы 2: "))
# Вводим начальные условия
E = start_table1[str(input_row1)][0]
J = start_table1[str(input_row1)][1]
alpha = start_table1[str(input_row1)][2]
omega = start_table2[str(input_row2)][0]
R = start_table2[str(input_row2)][1]
L = start_table2[str(input_row2)][2]
C = (start_table2[str(input_row2)][3])

table1_dictionary = {0: ["E", "J", "α", "ω", "R", "L", "C"],
                     1: ["B", "A", "град", "1/с", "Ом", "Гн", "мкФ"],
                     2: [str(E), str(J), str(alpha), str(omega), str(R), str(L), str(C)]}
table1 = doc.add_table(rows=3, cols=7, style='Table Grid')
for i in range(len(table1_dictionary)):
    for j in range(len(table1_dictionary[0])):
        table1.cell(i, j).text = str(table1_dictionary[i][j])
doc.add_picture('C:/Users/опасныйперчик/Desktop/MEGA/ТОЭ2/РГР4/Python/TOE4_1scheme/image/1.PNG')
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Рисунок 1.  Заданная схема электрической цепи").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
C= (start_table2[str(input_row2)][3])*(10**(-6))

# Функция создания жирного текста по центру
def bold_text(text):
    paragraph = doc.add_paragraph(str(text))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.runs[0]
    run.font.bold = True

bold_text("Цепь первого порядка")
doc.add_paragraph("При постоянном источнике тока e(t) = E после срабатывания ключа К1, "
                  "когда ключ К2 ещё не сработал, определяем напряжение i(t).").paragraph_format.first_line_indent = Mm(12.5)
doc.add_paragraph("1.1. Используем упрощённый классический метод, "
                  "когда дифференциальное уравнение для искомой функции  не составляется.").paragraph_format.first_line_indent = Mm(12.5)
text = doc.add_paragraph("1.1.1. Определяем независимые начальные условия (ННУ) при t=0: ")
text.paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(text,"i_{L}(0-).")
run = text.add_run("Схема до коммутации: установившийся режим, постоянный источник, С – разрыв, L – закоротка.")
text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Функция добавления картинки
def add_image(number, text):
    doc.add_picture(r'C:/Users/опасныйперчик/Desktop/MEGA/ТОЭ2/РГР4/Python/TOE4_1scheme/image/{}.PNG'.format(number))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(str(text)).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

add_image(2,"Рисунок 2.  ННУ")
math2docx.add_math(doc.add_paragraph(),  r"i_{L}(0-)=0\ A")

# Page №4

math2docx.add_math(doc.add_paragraph(),  r"i(0-)=\frac{E}{R}="+str(round((E/R),2))+r"\ A")
doc.add_paragraph("1.1.2. Определяем ЗНУ при t=0+. (Cхема после коммутации ключа К1). ").paragraph_format.first_line_indent = Mm(12.5)
add_image(3,"Рисунок 3.  ЗНУ")
doc.add_paragraph("По закону Ома:")
math2docx.add_math(doc.add_paragraph(),  r"i(0+)=\frac{E}{2R}=\frac{"+str(E)+"}{"+str(2*R)+"}="+str(round((E/(2*R)),2))+r"\ A")
text = doc.add_paragraph("1.1.3. Определяем принуждённую составляющую ")
text.paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(text,"i_{пр}(t) ")
text.add_run("при t=∞  (схема после коммутации ключа К1, установившейся режим, постоянный источник, С – разрыв, L – закоротка)")
text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

add_image(3,"Рисунок 3.  ПС")
doc.add_paragraph("По закону Ома:")
math2docx.add_math(doc.add_paragraph(),  r"i_{пр}=\frac{E}{R+\frac{R}{2}}=\frac{"+str(E)+"}{"+str(R)+r"+\frac{"+str(R)+"}{2}}="+str(round((E/(R+R/2)),2))+r"\ A")
text = doc.add_paragraph("1.1.4. Определяем корень характеристического уравнения p. Используем метод сопротивления цепи после коммутации (C → 1/Cp; L → Lp),  "
                  "причем ")
text.paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(text,r"R_{J}=∞,\ R_{E}=0.")
text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Page №5

add_image(4,"Рисунок 4. Схема определения корня характеристического уравнения")
math2docx.add_math(doc.add_paragraph(),  r"Z(p)=Lp+\frac{3R}{2}=0")
math2docx.add_math(doc.add_paragraph(),  r"p=-\frac{3R}{2L}=\frac{"+str(3*R)+"}{"+str(2*L)+"}="+str(round((-(3*R)/(2*L)),2))+r"\ 1/c")
doc.add_paragraph("1.1.5. Определяем постоянную интегрирования B:")
math2docx.add_math(doc.add_paragraph(),  r"B=i(0+)-i_{пр}="+str(round((E/(2*R)),2))+"-"+str(round((E/(R+R/2)),2))+"="+str(round(((E/(2*R))-(E/(R+R/2))),2))+r"\ B")
doc.add_paragraph("1.1.6. Окончательный результат:")
math2docx.add_math(doc.add_paragraph(),  r"i(t)=i_{пр}+Be^{pt}="+str(round((E/(R+R/2)),2))+"+("+str(round(((E/(2*R))-(E/(R+R/2))),2))+r")e^{"+str(round((-(3*R)/(2*L)),2))+r"t}\ ,B")
math2docx.add_math(doc.add_paragraph(),  r"\tau =\frac{1}{|p|}=\frac{1}{|"+str(round((-(3*R)/(2*L)),2))+"|}="+str(round(1/(abs(-(3*R)/(2*L))),4))+r"\ c - постоянная\ времени.")

# Начало и конец изменения значения X
t = np.linspace(-0.01,
                    0.1,
                    100)

#Функция создания и добавления графика
def create_plot(name,ylabel,text,limit_up_y,function_t,limit_down_y = 0):
    plt.clf()
    memfile = BytesIO()
    plt.title(str(name))
    plt.xlabel('t,c')
    plt.ylabel(str(ylabel))
    plt.xlim(0, 0.1)
    plt.ylim(int(limit_down_y), int(limit_up_y))
    plt.grid(True)
    y = function_t
    plt.plot(t, y)
    plt.savefig(memfile)
    doc.add_picture(memfile, width=Inches(5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    memfile.close()
    doc.add_paragraph(str(text)).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

create_plot("Переходный процесс",
            "i, A",
"Рисунок 5. Классический метод",
            2,
            (E/(R+R/2))+((E/(2*R))-(E/(R+R/2)))*np.exp((-(3*R*t)/(2*L))))

doc.add_paragraph("1.2. Используем операторный метод.").paragraph_format.first_line_indent = Mm(12.5)
doc.add_paragraph("1.2.1. Находим независимые начальные условия (п. 1.1.1)").paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(doc.add_paragraph(),  r"i_{L}(0-)=0\ A")
doc.add_paragraph("1.2.2. В операторной схеме после коммутации используем метод контурных токов:").paragraph_format.first_line_indent = Mm(12.5)
add_image(5,"Рисунок 6. Операторная схема")
math2docx.add_math(doc.add_paragraph(),r"\begin{displaymath}"
r"\left\{ \begin{array}{ll}"
r"I_{11}(p)\cdot 2R+I_{22}(p)\cdot R=\frac{E}{p}  \\"
r"I_{11}(p)\cdot R+I_{22}(p)\cdot (2R+Lp)=-L\cdot i_{L}(0-)=0 "
r"\end{array} \right."
r"\end{displaymath}")
doc.add_paragraph("Определяем операторное изображение искомого тока:").paragraph_format.first_line_indent = Mm(12.5)

#Решение уравнения
x, y,p = sp.symbols('x y p')
eq1 = sp.Eq(2*R*x + R*y, E/p)
eq2 = sp.Eq(R*x + y*(2*R+L*p), 0)
solution = sp.solve((eq1, eq2), (x, y))

math2docx.add_math(doc.add_paragraph(),r"i(p)=I_{11}(p)="+str(solution[x])+"")
doc.add_paragraph("По теореме разложения находим i(t):").paragraph_format.first_line_indent = Mm(12.5)

#Выделяем знаменатель для нахождения его нулей
denominator = re.findall(r'\((.*?)\)', str(solution[x]))[1]
match = re.search(r'(\d*\.?\d*)\*p\*\*2', denominator)
if match:
    number_in_square = match.group(1)
else:
    number_in_square = str(1)
match = re.findall(r'(\d*\.?\d*)\*p$', denominator)[0]
if match:
    number = match
else:
    number = str(1)

#Решатель уравнений
def solver(equation, value):
    p = value
    return eval(equation.replace("p", str(value)))

math2docx.add_math(doc.add_paragraph(),str(number_in_square)+"p^{2}+"+str(number)+"p=0")
numer, denom = sp.fraction(solution[x])
zeros_from_denom = sp.solve(denom)
derivative_of_denom = str(sp.Derivative(denom,p).doit())
math2docx.add_math(doc.add_paragraph(),"p_{0}="+str(round((-(3*R)/(2*L)),2))+",p_{1}="+str(round((zeros_from_denom[1]),2))+r"\quad ,1/c")
math2docx.add_math(doc.add_paragraph(),r"i(t)=\frac{"+str(solver((re.findall(r'\((.*?)\)', str(solution[x]))[0]),round((zeros_from_denom[1]),2)))+"}{"+str(solver(derivative_of_denom,round((zeros_from_denom[1]),2)))+"}e^{"+str(round((zeros_from_denom[1]),1))+r"}+"
                                            r"\frac{"+str(solver((re.findall(r'\((.*?)\)', str(solution[x]))[0]),round((zeros_from_denom[0]),0)))+"}{"+str(solver(derivative_of_denom,round((zeros_from_denom[0]),0)))+"}e^{"+str(round((-(3*R)/(2*L)),2))+r"\cdot t}="+str(round((E/(R+R/2)),2))+"+("+str(round(((E/(2*R))-(E/(R+R/2))),2))+r")e^{"+str(round((-(3*R)/(2*L)),2))+r"t}\ ,B")
create_plot("Переходный процесс",
            "i, A",
"Рисунок 7. Операторный метод",
            2,
            (E/(R+R/2))+((E/(2*R))-(E/(R+R/2)))*np.exp((-(3*R*t)/(2*L))))

text = doc.add_paragraph("2. При гармоническом источнике ЭДС ")
text.paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(text,r"e(t)=\sqrt{2}E\sin (\omega t+\alpha)")
text.add_run(" определить ток i(t).")
text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
text = doc.add_paragraph("2.1. Используем упрощённый классический метод, когда дифференциальное уравнение для искомой функции i(t)  не составляется.").paragraph_format.first_line_indent = Mm(12.5)
text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
text = doc.add_paragraph("2.1.1. ННУ. Определяем независимые начальные условия при  t=0-; ")
text.paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(text,r"i_{L}(0-)")
text.add_run(" (схема до коммутации установившийся режим, гармонический источник, символический метод). ")
text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

add_image(6,"Рисунок 8. ННУ")

#Переопределение Е
E = E*(cmath.cos(np.radians(alpha))+1j*cmath.sin(np.radians(alpha)))

complex_number = E/R
doc.add_paragraph("Ток течет через закорокту. ").paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(doc.add_paragraph(),r"\underline{E}="+str(abs(E))+"e^{j"+str(alpha)+r"} ,B")
math2docx.add_math(doc.add_paragraph(),r"X_{L}=\omega L="+str(L*omega)+r" ,Ом")
math2docx.add_math(doc.add_paragraph(),r"\underline{I}=\frac{\underline{E}}{R}="+str(round(complex_number.real, 2) + round(complex_number.imag, 2) * 1j)+r"="+str(round((abs(E/R)),2))+"e^{"+str(round((np.degrees(cmath.phase(E/R))),1))+"j} ,A")
math2docx.add_math(doc.add_paragraph(),r"i(t)="+str(round((abs(E/R)),2))+r"\sqrt{2} \sin ("+str(omega)+"t+("+str(round((np.degrees(cmath.phase(E/R))),1))+")) ,A;")
math2docx.add_math(doc.add_paragraph(),r"i(0)="+str(round((abs(E/R)),2))+r"\sqrt{2} \sin ("+str(round((np.degrees(cmath.phase(E/R))),1))+")="+str(round((((abs(E/R))*np.sqrt(2)*np.sin(cmath.phase(E/R)))),2))+" ,A;i_{L}(0)=0.")
doc.add_paragraph("2.1.2. Определяем ЗНУ при t=0+  (схема после коммутации ключа К1):  ").paragraph_format.first_line_indent = Mm(12.5)
add_image(7,"Рисунок 9. ЗНУ")
doc.add_paragraph("По Закону Ома:  ").paragraph_format.first_line_indent = Mm(12.5)

math2docx.add_math(doc.add_paragraph(),r"e(t)="+str(round((abs(E)),2))+r"\sqrt{2} \sin ("+str(omega)+"t+("+str(alpha)+")) ,B;")
math2docx.add_math(doc.add_paragraph(),r"e(0)="+str(round((abs(E)),2))+r"\sqrt{2} \sin ("+str(alpha)+")="+str(round((((abs(E))*np.sqrt(2)*np.sin(np.radians(alpha)))),2))+" ,B;")
math2docx.add_math(doc.add_paragraph(),r"i(0+)=\frac{E(0)}{2R}="+str(round((((((abs(E))*np.sqrt(2)*np.sin(np.radians(alpha)))))/(2*R)),2))+",A")
text = doc.add_paragraph("2.1.3. Определяем принуждённую составляющую  ")
text.paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(text,r"i_{пр}(0-)")
text.add_run(" при t=∞ (схема после коммутации ключа К1: установившейся режим, гармонический источник, символический метод):  ")
text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
add_image(8,"Рисунок 10. ПС")
doc.add_paragraph("По Закону Ома:  ").paragraph_format.first_line_indent = Mm(12.5)
complex_number = R+(R*(R+1j*(omega*L)))/(2*R+1j*(omega*L))
math2docx.add_math(doc.add_paragraph(),r"\underline{Z}=R+\frac{R\cdot (R+jX_{L})}{R+R+jX_{L}}="+str(round(complex_number.real, 2) + round(complex_number.imag, 2) * 1j)+r"="+str(round((abs(complex_number)),2))+"e^{"+str(round((np.degrees(cmath.phase(complex_number))),1))+"j} ,Ом")
complex_number = E/(R+(R*(R+1j*(omega*L)))/(2*R+1j*(omega*L)))
math2docx.add_math(doc.add_paragraph(),r"\underline{I_{пр}}=\frac{\underline{E}}{\underline{Z}}="+str(round(complex_number.real, 2) + round(complex_number.imag, 2) * 1j)+r"="+str(round((abs(complex_number)),2))+"e^{"+str(round((np.degrees(cmath.phase(complex_number))),1))+"j} ,A")
math2docx.add_math(doc.add_paragraph(),r"i_{пр}(t)="+str(round((abs(complex_number)),2))+r"\sqrt{2} \sin ("+str(omega)+"t+("+str(round((np.degrees(cmath.phase(complex_number))),1))+")) ,A;")
math2docx.add_math(doc.add_paragraph(),r"i_{пр}(0)="+str(round((abs(complex_number)),2))+r"\sqrt{2} \sin ("+str(round((cmath.phase(complex_number)),1))+")="+str(round((((abs(complex_number))*np.sqrt(2)*np.sin(cmath.phase(complex_number)))),2))+" ,A;")
doc.add_paragraph("2.1.4. Определяем корень характеристического уравнения : Используем метод сопротивления цепи после коммутации. Аналогично п. 1.1.4 получаем p="+str(round((-(3*R)/(2*L)),2))+" ,1/c").paragraph_format.first_line_indent = Mm(12.5)
doc.add_paragraph("2.1.5. Определяем постоянную интегрирования B:").paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(doc.add_paragraph(),r"B=i(0+)-i_{пр}(0)="+str(round(((((((abs(E))*np.sqrt(2)*np.sin(np.radians(alpha)))))/(2*R))-(((abs(complex_number))*np.sqrt(2)*np.sin(cmath.phase(complex_number))))),2))+",A;")
doc.add_paragraph("2.1.6. Окончательный результат").paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(doc.add_paragraph(),  r"i(t)=i_{пр}+Be^{pt}="+str(round((abs(complex_number)),2))+r"\sqrt{2} \sin ("+str(omega)+"t+("+str(round((np.degrees(cmath.phase(complex_number))),1))+"))+("
                   +str(round(((((((abs(E))*np.sqrt(2)*np.sin(np.radians(alpha)))))/(2*R))-(((abs(complex_number))*np.sqrt(2)*np.sin(cmath.phase(complex_number))))),2))+")e^{"+str(round((-(3*R)/(2*L)),2))+"t} ,A")
text = doc.add_paragraph("причем  ")
text.paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(text,r"\tau = \frac{1}{|p|}="+str((round((1/(zeros_from_denom[0])),4)))+r",\frac{1}{c};t_{п}=5\tau="+str(round((5*(1/(-(3*R)/(2*L)))),2))+r",c;"
                                                                                                    r"T=\frac{2\pi}{\omega}="+str(round((((2*np.pi)/(omega))),4))+"c,")
text.add_run(" -период принужденной составляющей.  ")
text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
create_plot("Переходный процесс",
            "i, A",
"Рисунок 11. Классичекий метод с гармоническим источником",
            2,
            (((abs(complex_number))*np.sqrt(2)*np.sin(omega*t+(cmath.phase(complex_number)))+((((abs(E))*np.sqrt(2)*np.sin(np.radians(alpha)))/(2*R))-((abs(complex_number))*np.sqrt(2)*np.sin(cmath.phase(complex_number))))*np.exp((-(3*R*t)/(2*L))))),
                   -2)
doc.add_paragraph("2.2. Используем комбинированный операторно-классический метод для определения i(t). ").paragraph_format.first_line_indent = Mm(12.5)
doc.add_paragraph("2.2.1. Находим независимые начальные условия (п. 2.1.1): ").paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(doc.add_paragraph(),r"i_{L}(0)=0")
text = doc.add_paragraph("2.2.2. Определяем принуждённые составляющие  ")
text.paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(text,r"i_{пр}(t),i_{Lпр}(t)")
text.add_run(" при t=∞ (схема после коммутации ключа К1: установившийся режим, гармонический источник, символический метод.) ")
text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
add_image(10,"Рисунок 12. ПС")
doc.add_paragraph("По закону Ома:").paragraph_format.first_line_indent = Mm(12.5)
complex_number = R+(R*(R+1j*(omega*L)))/(2*R+1j*(omega*L))
math2docx.add_math(doc.add_paragraph(),r"\underline{Z}=R+\frac{R\cdot (R+jX_{L})}{R+R+jX_{L}}="+str(round(complex_number.real, 2) + round(complex_number.imag, 2) * 1j)+r"="+str(round((abs(complex_number)),2))+"e^{"+str(round((np.degrees(cmath.phase(complex_number))),1))+"j} ,Ом")
complex_number = E/(R+(R*(R+1j*(omega*L)))/(2*R+1j*(omega*L)))
math2docx.add_math(doc.add_paragraph(),r"\underline{I_{пр}}=\frac{\underline{E}}{\underline{Z}}="+str(round(complex_number.real, 2) + round(complex_number.imag, 2) * 1j)+r"="+str(round((abs(complex_number)),2))+"e^{"+str(round((np.degrees(cmath.phase(complex_number))),1))+"j} ,A")
math2docx.add_math(doc.add_paragraph(),r"i_{пр}(t)="+str(round((abs(complex_number)),2))+r"\sqrt{2} \sin ("+str(omega)+"t+("+str(round((np.degrees(cmath.phase(complex_number))),1))+")) ,A;")
math2docx.add_math(doc.add_paragraph(),r"i_{пр}(0)="+str(round((abs(complex_number)),2))+r"\sqrt{2} \sin ("+str(round((cmath.phase(complex_number)),1))+")="+str(round((((abs(complex_number))*np.sqrt(2)*np.sin(cmath.phase(complex_number)))),2))+" ,A;")
doc.add_paragraph("По правилу разброса:").paragraph_format.first_line_indent = Mm(12.5)
complex_number = (E/(R+(R*(R+1j*(omega*L)))/(2*R+1j*(omega*L))))*(R/(2*R+1j*(omega*L)))
math2docx.add_math(doc.add_paragraph(),r"\underline{I_{Lпр}}=I_{пр}\cdot \frac{R}{2R+jX_{L}}="+str(round(complex_number.real, 2) + round(complex_number.imag, 2) * 1j)+r"="+str(round((abs(complex_number)),2))+"e^{"+str(round((np.degrees(cmath.phase(complex_number))),1))+"j} ,A")
math2docx.add_math(doc.add_paragraph(),r"i_{Lпр}(t)="+str(round((abs(complex_number)),2))+r"\sqrt{2} \sin ("+str(omega)+"t+("+str(round((np.degrees(cmath.phase(complex_number))),1))+")) ,A;")
math2docx.add_math(doc.add_paragraph(),r"i_{Lпр}(0)="+str(round((abs(complex_number)),2))+r"\sqrt{2} \sin ("+str(round((cmath.phase(complex_number)),1))+")="+str(round((((abs(complex_number))*np.sqrt(2)*np.sin(cmath.phase(complex_number)))),2))+" ,A;")
doc.add_paragraph("2.2.3. Определяем начальное значение свободной составляющей тока на индуктивности").paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(doc.add_paragraph(),r"i_{Lсв}(0)=i_{L}(0)-i_{Lпр}(0)="+str(round((-((abs(complex_number))*np.sqrt(2)*np.sin(cmath.phase(complex_number)))),3))+" ,A;")
doc.add_paragraph("2.2.4. Рассчитываем операторную схему замещения для свободных составляющих.").paragraph_format.first_line_indent = Mm(12.5)
add_image(11,"Рисунок 13. Операторная схема")
doc.add_paragraph("По закону Ома и по теореме разложения находим i(t):").paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(doc.add_paragraph(),r"I(p)=\frac{L\cdot i_{Lсв}(0)}{pL+1.5R} \cdot \frac{R}{2R};")
p = sp.symbols('p')
math2docx.add_math(doc.add_paragraph(),r"I(p)="+str(sp.simplify((0.5*L*(round((-((abs(complex_number))*np.sqrt(2)*np.sin(cmath.phase(complex_number)))),3)))/(1.5*R+L*p)))+r"=\frac{D(p)}{B(p)};")
complex_number = E/(R+(R*(R+1j*(omega*L)))/(2*R+1j*(omega*L)))
math2docx.add_math(doc.add_paragraph(),
r"\begin{array}{ll}"
r"I(p)=i_{пр}(t)+i_{св}(t)="+str(round((abs(complex_number)),2))+r"\sqrt{2} \sin ("+str(omega)+"t+("+str(round((np.degrees(cmath.phase(complex_number))),1))+r"))+\frac{D_{k}(p_{k})}{B'_{k}(p_{k})}e^{p_{k}t}= \\"
                   +str(round((abs(complex_number)), 2)) + r"\sqrt{2} \sin (" + str(omega) + "t+(" + str(round((np.degrees(cmath.phase(complex_number))), 1)) + "))+("
                   + str(round(((((((abs(E)) * np.sqrt(2) * np.sin(np.radians(alpha))))) / (2 * R)) - (((abs(complex_number)) * np.sqrt(2) * np.sin(cmath.phase(complex_number))))), 2)) + ")e^{" + str(round((-(3 * R) / (2 * L)), 2)) +r"t},A"
r"\end{array}")

create_plot("Переходный процесс",
            "i, A",
"Рисунок 14. Комбинированный метод",
            2,
            (((abs(complex_number))*np.sqrt(2)*np.sin(omega*t+(cmath.phase(complex_number)))+((((abs(E))*np.sqrt(2)*np.sin(np.radians(alpha)))/(2*R))-((abs(complex_number))*np.sqrt(2)*np.sin(cmath.phase(complex_number))))*np.exp((-(3*R*t)/(2*L))))),
                   -2)

#Пункт 3. Импульсный источник

text = doc.add_paragraph("3. При импульсном источнике напряжения  ")
text.paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(text,r"E(t)=E \cdot e^{2pt}="+str(abs(E))+"e^{"+str(round((-(2*3*R)/(2*L)),2))+r"t},B")
text.add_run(" (p – корень характеристического уравнения) и нулевых начальных условиях (ключ К1 сработал) определяем интегралом Дюамеля ток i(t).")
text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
doc.add_paragraph("3.1. Находим переходную характеристику h(t) для i(t) операторным методом").paragraph_format.first_line_indent = Mm(12.5)
doc.add_paragraph("Переходная характеристика системы h(p)   это реакция цепи в виде тока или напряжения на единичную возмущающую функцию I(t) источника, при нулевых начальных условиях").paragraph_format.first_line_indent = Mm(12.5)
add_image(12,"Рисунок 15. Операторная схема")
math2docx.add_math(doc.add_paragraph(),r"i_{L}(0)=0")

#Разбор полученной функции Дюамеля
p, t = sp.symbols('p t')
image_of_Duamel = sp.apart(sp.factor((1/(R+(1/((1/R)+1/(R+p*L)))))/p)).evalf(3)
denominator_1 = re.search(r'\((.*?)\)', str(image_of_Duamel)).group(1)
numerator_p = re.search(r'(-?\d+\.\d+)\/', str(image_of_Duamel)).group(1)
numerator_e = float(re.findall(r'(.*?)\/', str(image_of_Duamel))[0])
numerator_e = "{:f}".format(numerator_e)
original_of_Duamel = sp.inverse_laplace_transform(image_of_Duamel, p, t)
digit_exp = re.search(r'([-+] \d+\.\d+)', str(original_of_Duamel)).group(1).replace(" ","")
digit = re.findall(r'\d*\.\d+', str(original_of_Duamel))[0].replace(" ","")
math2docx.add_math(doc.add_paragraph(),r"h(p)=\frac{1}{p}\left(R+\left(\frac{1}{R}+\frac{1}{R+pL}\right)^{-1}\right)^{-1}=\frac{"+str(numerator_e)+r"}{"+str(denominator_1)+r"}+\frac{"+str(numerator_p)+r"}{p}=\frac{D(p)}{B(p)}")
math2docx.add_math(doc.add_paragraph(),"p="+str(round((-(3*R)/(2*L)),2))+",1/c")
math2docx.add_math(doc.add_paragraph(),r"h(t)="+str(digit)+"+("+str(digit_exp)+r")\cdot e^{"+str(round((-(3*R)/(2*L)),2))+r"t}")

doc.add_paragraph("Проверка:").paragraph_format.first_line_indent = Mm(12.5)
text = doc.add_paragraph("а)")
text.paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(text,r"t=0, h(0)="+str(round((1/(2*R)),3))+r"=\frac{1}{2R}=R_э(0)")
text.add_run(" -верно, так как L-разрыв;")
text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
text = doc.add_paragraph("б)")
text.paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(text,r"t=∞, h(∞)="+str(round((1/(1.5*R)),5))+r"=\frac{1}{R+0.5R}=R_э(∞)")
text.add_run(" -верно, так как L-закоротка;")
text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
doc.add_paragraph("3.2. Рассчитаем интегралом Дюамеля i(t):").paragraph_format.first_line_indent = Mm(12.5)

#Интеграл Дюамеля

derivative_E = str(sp.Derivative((abs(E)*sp.exp((-(2*3*R*p)/(2*L)))),p).doit())
derivative_E_digit = re.findall(r'([+-]\d+)', str(derivative_E))[0]
x, t = sp.symbols('x t')
integrator_1 =sp.integrate(((float(derivative_E_digit)*sp.exp((-(2*3*R*x)/(2*L))))
                   *float(digit))
                               ,(x,0,t))
integrator_2 = sp.integrate(((float(derivative_E_digit)*sp.exp((-(2*3*R*x)/(2*L))))
                   *float(digit_exp)*sp.exp((-(3*R*t)/(2*L)))*sp.exp(((3*R*x)/(2*L))))
                               ,(x,0,t))
duamel = sp.simplify(abs(E)*(float(digit)+float(digit_exp)*sp.exp((-(3*R*t)/(2*L))))+
                   integrator_1+integrator_2)
digit_small_e = float(re.findall(r'(\d*\.\d+)\*e', str(duamel))[0])
digit_big_e = float(re.findall(r'(\d*\.\d+)\*e', str(duamel))[1])
math2docx.add_math(doc.add_paragraph(),r"i_e(t)=E(0)h(t)+\int_0^t E^{'}(\tau)h(t-\tau )d\tau")
math2docx.add_math(doc.add_paragraph(),r"E(0)="+str(abs(E))+r",B")
math2docx.add_math(doc.add_paragraph(),r"E^{'}(\tau)="+str(derivative_E_digit)+r"e^{"+str(round((-(2*3*R)/(2*L)),2))+r"\tau}")
math2docx.add_math(doc.add_paragraph(),r"h(t-\tau)="+str(digit)+"+("+str(digit_exp)+r")\cdot e^{"+str(round((-(3*R)/(2*L)),2))+r"t}e^{"+str(round((-(3*R)/(2*L)),2))+r"\tau}")
math2docx.add_math(doc.add_paragraph(),r"i(t)=E(0)h(t)+\int_0^t E^{'}(\tau)h(t-\tau )d\tau=")
math2docx.add_math(doc.add_paragraph(),r"="+str(abs(E))+r"\cdot ("+str(digit)+"+("+str(digit_exp)+r")\cdot e^{"+str(round((-(3*R)/(2*L)),2))+r"t})+")
math2docx.add_math(doc.add_paragraph(),r"+\int_0^t "+str(derivative_E_digit)+r"e^{"+str(round((-(2*3*R)/(2*L)),2))+r"\tau} \cdot ("+str(digit)+"+("+str(digit_exp)+r")\cdot e^{"+str(round((-(3*R)/(2*L)),2))+r"t}e^{"+str(round((-(3*R)/(2*L)),2))+r"\tau})d \tau =")
math2docx.add_math(doc.add_paragraph(),r"="+str(digit_small_e)+r"e^{"+str(round((-(3*R)/(2*L)),2))+r"t}+"+str(digit_big_e)+r"e^{"+str(round((-(2*3*R)/(2*L)),2))+r"t},A")


# Начало и конец изменения значения X
t = np.linspace(-0.01,
                    0.1,
                    100)

create_plot("Переходный процесс",
            "i, A",
"Рисунок 16. Импульсный источник",
            2,
            (float(digit_small_e)*np.exp(-(3*R*t)/(2*L))+float(digit_big_e)*np.exp(-(2*3*R*t)/(2*L))),
                   -2)

#Цепь второго порядка

#Переопределение Е
E = start_table1[str(input_row1)][0]

bold_text("Цепь второго порядка")
doc.add_paragraph("4. При постоянном источнике тока E(t) = E  после срабатывания ключа К2 определяем ток i(t)(Ключ К1 давно уже сработал).").paragraph_format.first_line_indent = Mm(12.5)
doc.add_paragraph("4.1. Используем упрощённый классический метод, когда дифференциальное уравнение для искомой функции i(t) не составляется.").paragraph_format.first_line_indent = Mm(12.5)
text = doc.add_paragraph("4.1.1. Определяем независимые начальные условия (ННУ): ")
text.paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(text,r"i_{L}(0-), u_{C}(0-),")
text.add_run(" при t=0- (Схема до коммутации: установившийся режим, постоянный источник, С – разрыв, L – закоротка). ")
text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
add_image(13,"Рисунок 17. ННУ")
text = doc.add_paragraph("Находим по правилу разброса")
text.paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(text,r"i_{L}(0-):")
text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
math2docx.add_math(doc.add_paragraph(),r"i_{L}(0-)=\frac{E}{R+0.5R} \cdot 0.5="+str(round((E/(R+0.5*R)),2))+",A")
math2docx.add_math(doc.add_paragraph(),r"u_{C}(0-)=0")
doc.add_paragraph("4.1.2. Определяем ЗНУ при t=0+ (Схема после коммутации ключа К2): ")
add_image(14,"Рисунок 18. ЗНУ")
math2docx.add_math(doc.add_paragraph(),r"i_{L}(0+)=i_{L}(0-)")
math2docx.add_math(doc.add_paragraph(),r"u_{C}(0+)=u_{C}(0-)")
doc.add_paragraph("По закону Ома:")
math2docx.add_math(doc.add_paragraph(),r"i_{C}(0+)=\frac{E}{R}-i_{L}(0+)="+str(round(((E/R)-(E/(R+0.5*R))),2))+",A")
text = doc.add_paragraph("4.1.3. Определяем принуждённую составляющую ")
text.paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(text,r"i_{Lпр},U_{Cпр} ")
text.add_run(" при  t=∞  (Схема после коммутации ключа К2: установившийся режим, постоянный источник, С – разрыв, L – закоротка); ")
text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
add_image(15,"Рисунок 19. ПС")
doc.add_paragraph("По закону Ома:")
math2docx.add_math(doc.add_paragraph(),r"i_{Lпр}=\frac{E}{1.5R} \cdot 0.5="+str(round(((E*0.5)/(1.5*R)),2))+",A")
math2docx.add_math(doc.add_paragraph(),r"U_{Cпр}=i_{Lпр} \cdot R="+str(round(((E*0.5*R)/(1.5*R)),2))+",B")

text = doc.add_paragraph("4.1.4. Определяем корень характеристического уравнения p. Используем метод сопротивления цепи после коммутации (C → 1/Cp; L → Lp),  "
                  "причем ")
text.paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(text,r"R_{J}=∞,\ R_{E}=0.")
text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
add_image(16,"Рисунок 20. Корень характеристического уравнения")
math2docx.add_math(doc.add_paragraph(),r"Z(p)=0")
math2docx.add_math(doc.add_paragraph(),r"\frac{1}{Cp}+\frac{(Lp+R)\cdot 0.5R}{Lp+R+0.5R}=0")

p = sp.symbols('p')
roots_of_equation = sp.solve((1)/(C*p)+((L*p+R)*0.5*R)/(L*p+1.5*R))
real_digit = re.search(r'(\d*\.\d+) ', str(roots_of_equation[0])).group(1)
imag_digit = re.search(r'(\d*\.\d+)\*', str(roots_of_equation[0])).group(1)
math2docx.add_math(doc.add_paragraph(),r"p_{1}=-"+str(round((float(real_digit)),2))+"+"+str(round((float(imag_digit)),2))+"j, c^{-1};")
math2docx.add_math(doc.add_paragraph(),r"p_{2}=-"+str(round((float(real_digit)),2))+"-"+str(round((float(imag_digit)),2))+"j, c^{-1};")
doc.add_paragraph("4.1.5. Определяем постоянные интегрирования: т.к. характеристическое уравнение имеет 2 корня, то свободная составляющая будет иметь следующий вид:")
math2docx.add_math(doc.add_paragraph(),r"U_{Cсв}(t)=B_{1}e^{p_{1}t}+B_{2}e^{p_{2}t}")
doc.add_paragraph("а полное напряжение:")
math2docx.add_math(doc.add_paragraph(),r"U_{C}(t)=U_{Cпр}+U_{Cсв}(t)=U_{Cпр}+B_{1}e^{p_{1}t}+B_{2}e^{p_{2}t}")
doc.add_paragraph("Составляем систему и определяем значений постоянных интегрирования:")

math2docx.add_math(doc.add_paragraph(),r"\begin{displaymath}"
r"\left\{ \begin{array}{ll}"
r"B_{1}+B_{2}=U_{C}(0-)-U_{Cпр}  \\"
r"p_{1}B_{1}+p_{2}B_{2}=\frac{i_{C}(0+)}{C} "
r"\end{array} \right."
r"\end{displaymath}")
doc.add_paragraph("Решим систему уравнений и найдем постоянные интегрирования:").paragraph_format.first_line_indent = Mm(12.5)

#Решение уравнения
B1,B2 = sp.symbols('B1 B2')
eq1 = sp.Eq(B1+B2, 0-((E*0.5*R)/(1.5*R)))
eq2 = sp.Eq(B1*(-float(real_digit)+float(imag_digit)*1j)+B2*(-float(real_digit)-float(imag_digit)*1j),((E/R)-(E/(R+0.5*R)))/C)
solution = sp.solve((eq1, eq2), (B1, B2))

real_digit_B = re.search(r'(\d*\.\d+) ', str(solution[B1])).group(1)
imag_digit_B = re.search(r'(\d*\.\d+)\*', str(solution[B1])).group(1)
complex_B = -float(real_digit_B)+float(imag_digit_B)*1j
math2docx.add_math(doc.add_paragraph(),r"B_{1,2}=-"+str(round((float(real_digit_B)),2))+r"\pm j"+str(round((float(imag_digit_B)),2))+"="+str(round((abs(complex_B)),2))+r"e^{\pm j"+str(round((np.degrees(cmath.phase(complex_B))),1))+"}")

doc.add_paragraph("Таким образом, полное напряжение на конденсаторе:").paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(doc.add_paragraph(),r"U_{C}(t)=2B_{1} \cdot e^{-\delta t } \cdot \cos (\omega_{св} t+arg(B_{1}))+U_{Cпр}=")
math2docx.add_math(doc.add_paragraph(),"="+str(round((2*abs(complex_B)),2))+r"e^{-"+str(round((float(real_digit)),2))+r"t} \cdot \cos ("+str(round((float(imag_digit)),2))+"t+"+str(round((np.degrees(cmath.phase(complex_B))),1))+")+"+str(round(((E*0.5*R)/(1.5*R)),2))+",B")
doc.add_paragraph("4.1.6. Окончательный результат –").paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(doc.add_paragraph(),r"i(t)=\frac{E-U_{C}(t)}{R}="+str(round(((E-((E*0.5*R)/(1.5*R)))/R),2))+"-"+str(round((2*abs(complex_B)/R),2))+r"e^{-"+str(round((float(real_digit)),2))+r"t} \cdot \cos ("+str(round((float(imag_digit)),2))+"t+"+str(round((np.degrees(cmath.phase(complex_B))),1))+"),A")
text = doc.add_paragraph("причем  ")
text.paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(text,r"\tau = \frac{1}{|p|}="+str((round((1/float(real_digit)),4)))+r",\frac{1}{c};t_{п}=5\tau="+str(round((5/float(real_digit)),2))+r",c;"
                                                                                                    r"T=\frac{2\pi}{\omega}="+str(round(((2*np.pi/((float(imag_digit))))),4))+"c,")
text.add_run(" -период принужденной составляющей.  ")
text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
doc.add_paragraph("4.1.7. На интервале времени 0<t<tп строим график:").paragraph_format.first_line_indent = Mm(12.5)

create_plot("Переходный процесс",
            "i, A",
"Рисунок 21. График переходного процесса",
            6,
            (((E-((E*0.5*R)/(1.5*R)))/R)-(2*abs(complex_B)/R)*np.exp(-t*float(real_digit))*np.cos(float(imag_digit)*t+(cmath.phase(complex_B)))),
                   0)

#Операторный метод
doc.add_paragraph("4.2. Используем операторный метод для определения .").paragraph_format.first_line_indent = Mm(12.5)
doc.add_paragraph("4.2.1. Из расчёта установившегося режима до коммутации находим независимые начальные условия (п. 4.1.1):").paragraph_format.first_line_indent = Mm(12.5)
text = doc.add_paragraph("Находим по правилу разброса")
text.paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(text,r"i_{L}(0-):")
text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
math2docx.add_math(doc.add_paragraph(),r"i_{L}(0-)=\frac{E}{R+0.5R} \cdot 0.5="+str(round(((E*0.5)/(R+0.5*R)),2))+",A")
math2docx.add_math(doc.add_paragraph(),r"u_{C}(0-)=0")
doc.add_paragraph("4.2.2. В операторной схеме после коммутации используем метод узловых потенциалов:").paragraph_format.first_line_indent = Mm(12.5)
add_image(17,"Рисунок 22. Операторная схема")
math2docx.add_math(doc.add_paragraph(),r"\phi_{a}=0")
math2docx.add_math(doc.add_paragraph(),r"\phi_{b}(p)(\frac{1}{R} + \frac{1}{R+pL} + \frac{1}{R} +Cp)=\frac{E}{pR} - \frac{Li_{L}}{R+Lp}")

#Определяем потенциал
potential_b = sp.simplify((E/(p*R)-((L*(E/(R+0.5*R)))/(R+L*p)))/(1/R+1/(R+p*L)+1/R+C*p))
denominator = re.findall(r'\((.*?)\)', str(potential_b))[1]
match = re.search(r'(\d*\.?\d*)\*p\*\*2', denominator)
if match:
    number_p_square = round((float(match.group(1))),5)
else:
    number_p_square = str(1)
match = re.search(r'(\d*\.?\d*)\*p ', denominator)
if match:
    number_p = round((float(match.group(1))),5)
else:
    number_p = str(1)
match = re.search(r'(\d*\.?\d*)$', denominator)
if match:
    number = round((float(match.group(1))),5)
else:
    number = str(0)
numer, denom = sp.fraction(potential_b)
match = re.search(r'(\d*\.?\d*)\*', str(numer))
if match:
    numer_number_p = round((float(match.group(1))),5)
else:
    numer_number_p = str(1)
match = re.search(r'(\d*\.?\d*)$', str(numer))
if match:
    numer_number = round((float(match.group(1))),5)
else:
    numer_number = str(0)
math2docx.add_math(doc.add_paragraph(),r"\phi_{b}(p)=\frac{"+str(numer_number_p)+"p+"+str(numer_number)+"}{p("+str(number_p_square)+"p^{2}+"+str(number_p)+"p+"+str(number)+")}")
doc.add_paragraph("Используя  MathCAD, получим операторное изображение искомого тока:").paragraph_format.first_line_indent = Mm(12.5)

#Определяем ток
currrent_p = sp.apart(sp.simplify((-((E/(p*R)-((L*(E/(R+0.5*R)))/(R+L*p)))/(1/R+1/(R+p*L)+1/R+C*p))+E/p)/R)).evalf(3)
numerator = re.findall(r'(.*?)\/', str(currrent_p))[0]
numerator_p = re.search(r'(-?\d+\.\d+)\/', str(currrent_p)).group(1)
match = re.search(r'\/\((.*?)\*p\*\*2', str(currrent_p))
if match:
    denom_number_p2 = round((float(match.group(1))),5)
    denom_number_p2 = "{:f}".format(denom_number_p2)
else:
    denom_number_p2 = str(1)
match = re.search(r' (\d*\.?\d*)\*p ', str(currrent_p))
if match:
    denom_number_p = round((float(match.group(1))),5)
else:
    denom_number_p = str(1)
match = re.search(r' (\d*\.?\d*)\) ', str(currrent_p))
if match:
    denom_number = round((float(match.group(1))),5)
else:
    denom_number = str(0)
math2docx.add_math(doc.add_paragraph(),r"i(p)=\frac{"+str(numerator_p)+r"}{p} +\frac{"+str(numerator)+"}{"+str(denom_number_p2)+r"p^{2}+"+str(denom_number_p)+"p+"+str(denom_number)+r"}=\frac{D(p)}{B(p)}")
doc.add_paragraph("4.2.3. По теореме разложения находим искомый ток i(t) при помощи программы MathCAD:").paragraph_format.first_line_indent = Mm(12.5)
currrent_p = sp.simplify((-((E/(p*R)-((L*(E/(R+0.5*R)))/(R+L*p)))/(1/R+1/(R+p*L)+1/R+C*p))+E/p)/R).evalf(3)
numer, denom = sp.fraction(currrent_p)
zeros_from_denom = sp.solve(denom)
real_digit_zero = re.search(r'(\d*\.\d+) ', str(zeros_from_denom[1])).group(1)
imag_digit_zero = re.search(r'(\d*\.\d+)\*', str(zeros_from_denom[1])).group(1)
complex_zero_plus = -round((float(real_digit_zero)),2)+round((float(imag_digit_zero)),2)*1j
complex_zero_minus = -round((float(real_digit_zero)),2)-round((float(imag_digit_zero)),2)*1j
p = sp.symbols('p')
mini_denom = float(denom_number_p2)*p**2+float(denom_number_p)*p
derivative_of_denom = str(sp.Derivative(mini_denom, p).doit())
math2docx.add_math(doc.add_paragraph(),r"B(p)="+str(denom_number_p2)+r"p^{2}+"+str(denom_number_p)+"p+"+str(denom_number)+"=0")
math2docx.add_math(doc.add_paragraph(),r"p_{1}="+str(zeros_from_denom[0])+",p_{2}="+str(complex_zero_plus)+",p_{2}="+str(complex_zero_plus)+r"=-\delta \pm j\omega_{св},1/c;")
math2docx.add_math(doc.add_paragraph(),r"B^{'}(p)="+str(derivative_of_denom)+";")
math2docx.add_math(doc.add_paragraph(),r"i(t)=\displaystyle \sum_{k=1}^{3} \frac{D_{k}(p_{k})}{B^{'}_{k}(p_{k})}e^{p_{k}t}="+str(numerator_p)+r"+\displaystyle \sum_{k=2}^{3} \frac{D_{k}(p_{k})}{B^{'}_{k}(p_{k})}e^{p_{k}t}=")
math2docx.add_math(doc.add_paragraph(),r"="+str(numerator_p)+r"+\frac{"+str(numerator.replace("p", str(complex_zero_plus)))+"}{"+str(derivative_of_denom.replace("p", str(complex_zero_plus)))+"}e^{"+str(complex_zero_plus)+"t}")
math2docx.add_math(doc.add_paragraph(),r"+"+str(numerator_p)+r"+\frac{"+str(numerator.replace("p", str(complex_zero_minus)))+"}{"+str(derivative_of_denom.replace("p", str(complex_zero_minus)))+"}e^{"+str(complex_zero_minus)+"t}")
complex_number_4 = eval(str((solver(numerator,complex_zero_plus))/(solver(derivative_of_denom,complex_zero_plus))))
real_digit_number = re.search(r'(\d*\.\d+)[+-]', str(complex_number_4)).group(1)
imag_digit_number = re.search(r'(\d*\.\d+)j', str(complex_number_4)).group(1)
complex_number_plus = -round((float(real_digit_number)),3)+round((float(imag_digit_number)),3)*1j
complex_number_minus = -round((float(real_digit_number)),3)-round((float(imag_digit_number)),3)*1j
math2docx.add_math(doc.add_paragraph(),r"="+str(numerator_p)+r"+"+str(complex_number_plus)+"e^{"+str(complex_zero_plus)+"t}+"+str(complex_number_minus)+"e^{"+str(complex_zero_minus)+"t}=")
math2docx.add_math(doc.add_paragraph(),r"="+str(numerator_p)+r"+"+str(round((abs(complex_number_plus)),3))+r"\cdot 2 \cdot e^{"+str(round((float(real_digit)),2))+r"t}e^{30j}e^{90j} \left( \frac{e^{"+str(round((float(imag_digit)),2))+r"jt}-e^{-"+str(round((float(imag_digit)),2))+r"jt}}{2j} \right) =")
math2docx.add_math(doc.add_paragraph(),r"="+str(numerator_p)+r"+"+str(round((2*abs(complex_number_plus)),3))+r"\cdot e^{"+str(round((float(real_digit)),2))+r"t} \sin ("+str(round((float(imag_digit)),2))+r"t+90+30)="+str(numerator_p)+r"+"+str(round((2*abs(complex_number_plus)),3))+r"\cdot e^{"+str(round((float(real_digit)),2))+r"t} \cos ("+str(round((float(imag_digit)),2))+r"t+30)=")
math2docx.add_math(doc.add_paragraph(),r"="+str(numerator_p)+r"-"+str(round((2*abs(complex_number_plus)),3))+r"\cdot e^{"+str(round((float(real_digit)),2))+r"t} \cos ("+str(round((float(imag_digit)),2))+r"t+150),A")

create_plot("Переходный процесс",
            "i, A",
"Рисунок 23. График переходного процесса",
            6,
            (((E-((E*0.5*R)/(1.5*R)))/R)-(2*abs(complex_B)/R)*np.exp(-t*float(real_digit))*np.cos(float(imag_digit)*t+(cmath.phase(complex_B)))),
                   0)

#Метод переменных состояний

doc.add_paragraph("4.3. Методом переменных состояния находим i(t).").paragraph_format.first_line_indent = Mm(12.5)
text = doc.add_paragraph("4.3.1. Определяем начальные значения переменных состояния  ")
text.paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(text,r"i_{L}(0),U_{C}(0) ")
text.add_run(" из схемы до коммутации. ")
text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
add_image(18,"Рисунок 24. ННУ")
doc.add_paragraph("Напряжение на конденсаторе равно нулю (ключ разомкнут): ").paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(doc.add_paragraph(),r"U_{C}(0)=0,B")
doc.add_paragraph("По закону Ома: ").paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(doc.add_paragraph(),r"i_{0}(0-)=\frac{E}{R+0.5R} ="+str(round((E/(R+0.5*R)),2))+",A")
doc.add_paragraph("Ток через индуктивность по правилу разброса : ").paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(doc.add_paragraph(),r"i_{0}(0-)=\frac{i_{0}(0-)}{2} ="+str(round(((E*0.5)/(R+0.5*R)),2))+",A")
doc.add_paragraph("4.3.2. Изображаем схему после коммутации и составляем дифференциальные уравнения").paragraph_format.first_line_indent = Mm(12.5)
add_image(19,"Рисунок 25. Схема после коммутации")
doc.add_paragraph("Составим систему уравнений по законам Кирхгофа:").paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(doc.add_paragraph(),r"\begin{displaymath}"
r"\left\{ \begin{array}{ll}"
r"L\frac{di_{L}(t)}{dt}+i_{L} \cdot R-U_{C}=0 \\"
r"i(t) \cdot R+U_{C}=E \\"
r"i_{R}(t) \cdot R-U_{C}=0 \\"
r"i(t)=i_{L}(t)+i_{R}(t)+C\frac{dU_{C}(t)}{dt} "                                    
r"\end{array} \right."
r"\end{displaymath}")
doc.add_paragraph("Выразим производные и токи:").paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(doc.add_paragraph(),r"\begin{displaymath}"
r"\left\{ \begin{array}{ll}"
r"\frac{di_{L}(t)}{dt}=-\frac{R}{L}i_{L}(t)+\frac{1}{L}U_{C}+0  \\"
r"i(t)=-\frac{1}{R}U_{C}+\frac{1}{R}E \\"
r"i_{R}(t)=\frac{1}{R}U_{C} \\"
r"C\frac{dU_{C}(t)}{dt}=i(t)-i_{L}(t)-i_{R}(t) "                                    
r"\end{array} \right."
r"\end{displaymath}")
doc.add_paragraph("Подставим в последнее уравнение:").paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(doc.add_paragraph(),r"\begin{displaymath}"
r"\left\{ \begin{array}{ll}"
r"\frac{di_{L}(t)}{dt}=-\frac{R}{L}i_{L}(t)+\frac{1}{L}U_{C}+0  \\"
r"C\frac{dU_{C}(t)}{dt}=-\frac{1}{R}U_{C}+\frac{1}{R}E-i_{L}(t)-\frac{1}{R}U_{C} "                                    
r"\end{array} \right."
r"\end{displaymath}")
doc.add_paragraph("Разрешим систему уравнений относительно производных (представим в нормальной форме Коши):").paragraph_format.first_line_indent = Mm(12.5)
math2docx.add_math(doc.add_paragraph(),r"\begin{displaymath}"
r"\left\{ \begin{array}{ll}"
r"\frac{di_{L}(t)}{dt}=-\frac{R}{L}i_{L}(t)+\frac{1}{L}U_{C}+0  \\"
r"\frac{dU_{C}(t)}{dt}=-\frac{1}{C}i_{L}(t)-\frac{2}{RC}U_{C}+\frac{1}{RC}E "                                    
r"\end{array} \right."
r"\end{displaymath}")
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph("Вывод: ")
doc.add_paragraph("В ходе проделанной работы были произведены расчеты различными методами: классическим, операторным.  "
                  "Наиболее удобным метом является операторный метод. Момент коммутации определяет начало переходного процесса,"
                  " при этом различают время непосредственно перед коммутацией  t(0-) и сразу после коммутации t(0+). "
                  "Время переходного процесса обусловлено временем изменения энергии электрического и магнитного полей "
                  "накопителей. В цепях без накопителей энергии переходный процесс отсутствует: после срабатывания ключа "
                  "(t=0+) в цепи сразу же возникает установившийся режим. ")
doc.save("new.docx")

